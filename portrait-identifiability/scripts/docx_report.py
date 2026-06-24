# -*- coding: utf-8 -*-
"""DOCX 报告生成器 — 肖像权可识别性检测

生成格式规范的 .docx 报告，方便在 Word / WPS 中查看、批注和留档。
HTML 报告的 PDF 导出存在留白过多的问题，DOCX 可以精确控制页边距和排版。
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Emu

from common import data_url_for_image, risk_label, risk_score


_PAGE_WIDTH = Cm(21.0)
_PAGE_HEIGHT = Cm(29.7)
_MARGIN = Cm(1.5)

_RISK_COLORS: dict[str, str] = {
    "high": "DC2626",
    "medium": "D97706",
    "low_to_medium": "CA8A04",
    "low": "16A34A",
    "unable_to_determine": "78716C",
}

_BODY_FONT = "等线"
_BODY_FONT_EAST = "等线"
_HEADING_FONT = "黑体"


def _set_cell_shading(cell, color_hex: str) -> None:
    shading_elm = cell._tc.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading_elm.append(shading)


def _set_run_font(run, size_pt: float = 10.5, bold: bool = False, color_hex: str | None = None):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = _BODY_FONT
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), _BODY_FONT_EAST)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True
    run.font.name = _HEADING_FONT
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), _HEADING_FONT)
    p.space_after = Pt(8)
    p.space_before = Pt(16 if level <= 2 else 10)
    return p


def _add_body(doc: Document, text: str, bold: bool = False, size_pt: float = 10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size_pt=size_pt, bold=bold)
    p.space_after = Pt(4)
    return p


def _read_image_bytes(image_path: str, max_width_cm: float = 14.0, max_height_cm: float = 10.0):
    try:
        from PIL import Image
        img = Image.open(image_path)
        w, h = img.size
        ratio = min(max_width_cm / (w * 0.026458), max_height_cm / (h * 0.026458), 1.0)
        new_w = Cm(w * 0.026458 * ratio)
        new_h = Cm(h * 0.026458 * ratio)
        buf = io.BytesIO()
        fmt = img.format or "PNG"
        if fmt.upper() in ("JPEG", "JPG"):
            img.save(buf, format="JPEG", quality=85)
        else:
            img.save(buf, format="PNG")
        buf.seek(0)
        return buf, new_w, new_h
    except Exception:
        buf = io.BytesIO(open(image_path, "rb").read())
        return buf, Cm(max_width_cm), Cm(max_height_cm)


def _add_image_cell(cell, image_path: str, label: str, max_w: float = 7.0, max_h: float = 6.0):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    result = _read_image_bytes(image_path, max_w, max_h)
    if result:
        buf, w, h = result
        run = p.add_run()
        run.add_picture(buf, width=w, height=h)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(label)
    _set_run_font(run2, size_pt=9, color_hex="78716C")


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "D6D3D1",
        })
        borders.append(element)
    tblPr.append(borders)


def _set_cell_text(cell, text: str, bold: bool = False, size_pt: float = 10, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    _set_run_font(run, size_pt=size_pt, bold=bold)


def _collect_unable_reasons(results: list[dict[str, Any]], limit: int = 5) -> list[str]:
    reasons: list[str] = []
    for item in results:
        if item.get("risk_level") != "unable_to_determine":
            continue
        precheck = item.get("local_precheck", {})
        for reason in precheck.get("reliability_issues", []) or []:
            if reason and reason not in reasons:
                reasons.append(str(reason))
        for reason in item.get("limitations", []) or []:
            if reason and reason not in reasons:
                reasons.append(str(reason))
        ai = item.get("ai_visual_comparison", {})
        for reason in ai.get("limitations", []) or []:
            if reason and reason not in reasons:
                reasons.append(str(reason))
        for reason in item.get("basis", []) or []:
            text = str(reason)
            if any(kw in text for kw in ["质量", "未检测", "多张人脸", "无法可靠", "比对失败"]):
                if text not in reasons:
                    reasons.append(text)
        if len(reasons) >= limit:
            break
    return reasons[:limit]


def generate_docx_report(
    query_image: str,
    results: list[dict[str, Any]],
    output_dir: str | Path,
    multimodal_provider: str = "无",
    total_pairs: int = 0,
) -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width = _PAGE_WIDTH
    section.page_height = _PAGE_HEIGHT
    section.top_margin = _MARGIN
    section.bottom_margin = _MARGIN
    section.left_margin = _MARGIN
    section.right_margin = _MARGIN

    style = doc.styles["Normal"]
    style.font.name = _BODY_FONT
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), _BODY_FONT_EAST)

    query_name = Path(query_image).name
    query_path = str(Path(query_image).resolve())

    risk_counts: dict[str, int] = {}
    for r in results:
        rl = r.get("risk_level", "unable_to_determine")
        risk_counts[rl] = risk_counts.get(rl, 0) + 1

    highest_risk = "unable_to_determine"
    for rl in ["high", "medium", "low_to_medium", "low"]:
        if risk_counts.get(rl, 0) > 0:
            highest_risk = rl
            break

    # 标题
    _add_heading(doc, "肖像权可识别性撞脸排查报告", level=1)
    meta_text = f"查询图：{query_name}    |    比对数：{total_pairs}    |    多模态引擎：{multimodal_provider}"
    _add_body(doc, meta_text, size_pt=9)
    doc.add_paragraph()

    # 综述区
    overview_table = doc.add_table(rows=1, cols=2)
    overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(overview_table)

    left_cell = overview_table.rows[0].cells[0]
    left_cell.width = Cm(6.5)
    _add_image_cell(left_cell, query_path, f"查询图：{query_name}", max_w=5.5, max_h=5.5)

    right_cell = overview_table.rows[0].cells[1]
    right_cell.width = Cm(11.5)
    for p in right_cell.paragraphs:
        p.clear()

    p = right_cell.paragraphs[0]
    run = p.add_run("综合风险等级：")
    _set_run_font(run, size_pt=12, bold=True)
    risk_label_text = risk_label(highest_risk)
    risk_color = _RISK_COLORS.get(highest_risk, "78716C")
    run2 = p.add_run(risk_label_text)
    _set_run_font(run2, size_pt=14, bold=True, color_hex=risk_color)

    dist_parts = []
    for rl in ["high", "medium", "low_to_medium", "low", "unable_to_determine"]:
        cnt = risk_counts.get(rl, 0)
        if cnt > 0:
            dist_parts.append(f"{risk_label(rl)} {cnt} 个")
    if dist_parts:
        p2 = right_cell.add_paragraph()
        run3 = p2.add_run("风险分布：")
        _set_run_font(run3, size_pt=10, bold=True)
        run4 = p2.add_run("  ".join(dist_parts))
        _set_run_font(run4, size_pt=10, color_hex="57534E")

    verdict_map = {
        "high": "存在较高撞脸风险，建议人工复核。本结果不能替代司法鉴定、律师意见或法院判断。",
        "medium": "存在一定撞脸风险，建议人工复核。本结果不能替代司法鉴定、律师意见或法院判断。",
        "low_to_medium": "存在较低程度的撞脸风险提示，建议人工关注。",
        "low": "在当前比对范围内未发现高相似候选对象。但这不代表不存在其他肖像权风险。",
        "unable_to_determine": "部分或全部比对对象无法可靠判断，建议人工复核。",
    }
    verdict = verdict_map.get(highest_risk, verdict_map["unable_to_determine"])
    p3 = right_cell.add_paragraph()
    run5 = p3.add_run(verdict)
    _set_run_font(run5, size_pt=10, color_hex="44403C")

    if highest_risk == "unable_to_determine":
        reasons = _collect_unable_reasons(results)
        if reasons:
            p4 = right_cell.add_paragraph()
            run6 = p4.add_run("无法判断原因：")
            _set_run_font(run6, size_pt=9, bold=True, color_hex="78716C")
            for reason in reasons:
                p5 = right_cell.add_paragraph()
                run7 = p5.add_run(f"  {reason}")
                _set_run_font(run7, size_pt=9, color_hex="78716C")

    doc.add_paragraph()

    # 详细比对
    _add_heading(doc, "详细比对结果", level=2)

    for idx, item in enumerate(results):
        if idx > 0:
            doc.add_page_break()

        rank = idx + 1
        ref_path = item.get("image_b", "")
        ref_name = Path(ref_path).name if ref_path else "未知"
        pi = item.get("pair_info", {})
        source = pi.get("source", "unknown")
        risk_lvl = item.get("risk_level", "unable_to_determine")
        path_label = item.get("analysis_path", "-")
        sim = item.get("overall_similarity")
        sim_str = f"{sim:.2f}" if sim is not None else "-"

        precheck = item.get("local_precheck", {})
        faces_a = precheck.get("faces_a", "?")
        faces_b = precheck.get("faces_b", "?")
        quality_a = precheck.get("quality_a", {})
        quality_b = precheck.get("quality_b", {})
        qa_grade = quality_a.get("grade", "?") if isinstance(quality_a, dict) else "?"
        qb_grade = quality_b.get("grade", "?") if isinstance(quality_b, dict) else "?"
        face_sim = precheck.get("face_similarity")

        ai = item.get("ai_visual_comparison")
        basis_items = ai.get("basis", []) if ai else item.get("basis", [])
        fc = ai.get("feature_comparison") if ai else None
        fusion_note = ai.get("insightface_fusion_note", "") if ai else ""

        risk_badge_str = risk_label(risk_lvl)
        _add_heading(doc, f"#{rank}  {ref_name[:50]}  —  {risk_badge_str}", level=3)

        # 图片并排
        img_table = doc.add_table(rows=1, cols=2)
        img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(img_table)
        img_table.rows[0].cells[0].width = Cm(7.5)
        img_table.rows[0].cells[1].width = Cm(7.5)
        _add_image_cell(img_table.rows[0].cells[0], query_path, f"查询图：{query_name}", max_w=6.5, max_h=5.5)
        _add_image_cell(img_table.rows[0].cells[1], ref_path, f"候选图：{ref_name}", max_w=6.5, max_h=5.5)

        doc.add_paragraph()

        # 基本信息表
        info_rows = [
            ("来源", source),
            ("分析路径", path_label),
            ("整体相似度", sim_str),
            ("人脸数", f"A={faces_a} / B={faces_b}"),
            ("画质等级", f"A-{qa_grade} / B-{qb_grade}"),
        ]
        face_sim_row = None
        if face_sim is not None and str(path_label).upper() != "B":
            face_sim_row = ("InsightFace 余弦相似度", f"{face_sim:.4f}")

        all_rows = info_rows[:2] + ([face_sim_row] if face_sim_row else []) + info_rows[2:]
        num_rows = len(all_rows)

        info_table = doc.add_table(rows=num_rows, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(info_table)

        for i, (label, value) in enumerate(all_rows):
            if value is None:
                continue
            _set_cell_text(info_table.rows[i].cells[0], label, bold=True, size_pt=9.5)
            _set_cell_text(info_table.rows[i].cells[1], str(value), size_pt=9.5)
            info_table.rows[i].cells[0].width = Cm(5)
            info_table.rows[i].cells[1].width = Cm(11)
            _set_cell_shading(info_table.rows[i].cells[0], "F5F5F4")

        # 目标对象提示
        if isinstance(faces_a, int) and faces_a > 1:
            _add_body(doc, "目标对象提示：图片 A 检测到多张人脸，当前结果无法确认比对目标。建议先裁剪到单一目标人脸后重新检测。", size_pt=9)
        elif isinstance(faces_b, int) and faces_b > 1:
            _add_body(doc, "目标对象提示：图片 B 检测到多张人脸，当前结果无法确认比对目标。建议先裁剪到单一目标人脸后重新检测。", size_pt=9)

        if fusion_note:
            _add_body(doc, fusion_note, size_pt=9)

        # 多模态特征对比表
        if fc and isinstance(fc, dict):
            feature_keys = [k for k, v in fc.items() if v]
            if feature_keys:
                _add_body(doc, "多模态特征对比：", bold=True, size_pt=10)
                ft = doc.add_table(rows=len(feature_keys), cols=2)
                ft.alignment = WD_TABLE_ALIGNMENT.CENTER
                _set_table_borders(ft)
                feature_names = {
                    "face_shape": "脸型轮廓",
                    "facial_layout": "五官布局",
                    "eyes_brows": "眼眉特征",
                    "nose_mouth": "鼻口特征",
                    "hair_hairline": "发型发际线",
                    "distinctive_features": "显著特征",
                }
                for fi, fk in enumerate(feature_keys):
                    _set_cell_text(ft.rows[fi].cells[0], feature_names.get(fk, fk), bold=True, size_pt=9)
                    _set_cell_text(ft.rows[fi].cells[1], str(fc[fk]), size_pt=9)
                    ft.rows[fi].cells[0].width = Cm(4)
                    ft.rows[fi].cells[1].width = Cm(12)
                    _set_cell_shading(ft.rows[fi].cells[0], "F5F5F4")

        if basis_items:
            _add_body(doc, "主要依据：", bold=True, size_pt=10)
            for bi in basis_items:
                _add_body(doc, f"  {bi}", size_pt=9.5)

        limitations = item.get("limitations", []) or []
        if ai and ai.get("limitations"):
            limitations.extend(ai["limitations"])
        if limitations:
            _add_body(doc, "局限性说明：", bold=True, size_pt=10)
            for lim in limitations:
                _add_body(doc, f"  {lim}", size_pt=9)

    # 限制说明
    doc.add_page_break()
    _add_heading(doc, "限制说明", level=2)
    _add_body(doc, "本结果仅表示在当前比对范围内的撞脸风险排查，不代表不存在其他肖像权风险。")
    _add_body(doc, "比对结果可能受图片质量、角度、遮挡、风格化程度和候选图来源影响。")

    if any(str(item.get("analysis_path", "")).upper() == "B" for item in results):
        _add_body(doc, "风格化图像场景不采用本地人脸 embedding 作为主要判断依据。")

    if any(
        isinstance(item.get("local_precheck", {}).get("faces_a"), int) and item.get("local_precheck", {}).get("faces_a") > 1
        or isinstance(item.get("local_precheck", {}).get("faces_b"), int) and item.get("local_precheck", {}).get("faces_b") > 1
        for item in results
    ):
        _add_body(doc, "多张人脸场景需要先裁剪到单一目标人脸，或明确指定目标人物后重新检测。")

    _add_body(doc, "本结果不能替代司法鉴定、律师意见或法院判断。")

    return doc


def save_docx_report(
    query_image: str,
    results: list[dict[str, Any]],
    output_dir: str | Path,
    multimodal_provider: str = "无",
    total_pairs: int = 0,
) -> str:
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = generate_docx_report(
        query_image=query_image,
        results=results,
        output_dir=str(out_dir),
        multimodal_provider=multimodal_provider,
        total_pairs=total_pairs,
    )
    report_path = out_dir / "clearance-report.docx"
    doc.save(str(report_path))
    return str(report_path)
